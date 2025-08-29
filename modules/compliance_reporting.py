import openai
import json
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from config import OPENAI_API_KEY, EMAIL_CONFIG
from modules.db import get_db_connection
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import pandas as pd
from docx import Document

# Inizializza il client OpenAI
client = openai.OpenAI(api_key=OPENAI_API_KEY)

def generate_compliance_report(report_type, period_start, period_end):
    """
    Genera report di compliance automatici
    """
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Recupera i dati necessari per il report
    if report_type == "GDPR":
        # Recupera tutti i client con i loro dati
        cursor.execute("""
            SELECT id, name, company, email, created_at
            FROM clients
            WHERE created_at BETWEEN %s AND %s
        """, (period_start, period_end))
        
        client_data = cursor.fetchall()
        
        # Recupera storico modifiche
        cursor.execute("""
            SELECT table_name, action, timestamp, user_id
            FROM audit_log
            WHERE timestamp BETWEEN %s AND %s
            AND table_name IN ('clients', 'policies', 'risks')
        """, (period_start, period_end))
        
        audit_data = cursor.fetchall()
        
        report_content = {
            "client_data_summary": f"Totale client registrati: {len(client_data)}",
            "data_processing_activities": audit_data,
            "retention_policies": "Dati mantenuti per 365 giorni",
            "security_measures": "Crittografia AES-256, backup giornalieri"
        }
        
    elif report_type == "SOX":
        # Recupera dati finanziari
        cursor.execute("""
            SELECT 
                SUM(amount) as total_premiums,
                COUNT(*) as policies_count
            FROM premiums
            WHERE payment_date BETWEEN %s AND %s
        """, (period_start, period_end))
        
        financial_data = cursor.fetchone()
        
        report_content = {
            "financial_summary": financial_data,
            "controls_implementation": "Controlli finanziari implementati",
            "audit_trail": "Tracciamento completo delle transazioni"
        }
        
    elif report_type == "IVASS":
        # Recupera dati assicurativi
        cursor.execute("""
            SELECT 
                r.risk_type,
                COUNT(p.id) as policies_count,
                SUM(pr.amount) as total_premiums
            FROM policies p
            JOIN risks r ON p.risk_id = r.id
            JOIN premiums pr ON p.id = pr.policy_id
            WHERE p.start_date BETWEEN %s AND %s
            GROUP BY r.risk_type
        """, (period_start, period_end))
        
        insurance_data = cursor.fetchall()
        
        report_content = {
            "portfolio_analysis": insurance_data,
            "compliance_status": "Conforme alle direttive IVASS",
            "risk_distribution": "Distribuzione rischi monitorata"
        }
    
    conn.close()
    
    # Genera contenuto del report con AI
    prompt = f"""
    Genera un report di compliance {report_type} per il periodo {period_start} a {period_end}.
    
    Dati disponibili:
    {json.dumps(report_content, indent=2, default=str)}
    
    Il report deve includere:
    1. Intestazione ufficiale
    2. Riepilogo esecutivo
    3. Dettagli tecnici
    4. Conclusioni e raccomandazioni
    5. Firma digitale (simulata)
    
    Formato: JSON con campi: title, executive_summary, technical_details, conclusions, signature
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Sei un esperto di compliance normativa nel settore assicurativo. Rispondi in formato JSON valido con i seguenti campi: title, executive_summary, technical_details, conclusions, signature"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        # Debug: stampa il tipo e il contenuto della risposta
        print(f"Tipo di response.choices[0].message.content: {type(response.choices[0].message.content)}")
        print(f"Contenuto della risposta: {response.choices[0].message.content}")
        
        # Estrai e pulisci il contenuto JSON
        # Gestisci sia il caso stringa che il caso dizionario
        message_content = response.choices[0].message.content
        if isinstance(message_content, dict):
            # Se è un dizionario, cerca il contenuto testuale
            if 'text' in message_content:
                content = message_content['text'].strip()
            else:
                # Prova a convertire in stringa
                content = str(message_content).strip()
        else:
            # Se è già una stringa
            content = str(message_content).strip()
        
        print(f"Contenuto dopo la pulizia: {content}")
        
        # Rimuovi eventuali marcatori di codice
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        print(f"Contenuto dopo la rimozione dei marcatori: {content}")
        
        report = json.loads(content)
        
        # Salva il report nel database
        report_id = save_compliance_report(report_type, period_start, period_end, report)
        
        # Genera e salva i file in tutti i formati
        # Convertiamo i valori del report in stringhe per evitare errori nei generatori di file
        sanitized_report = {}
        for key, value in report.items():
            if isinstance(value, (dict, list)):
                sanitized_report[key] = str(value)
            else:
                sanitized_report[key] = str(value) if value is not None else ""
        
        pdf_path = generate_pdf_report(report_id, report_type, sanitized_report)
        excel_path = generate_excel_report(report_id, report_type, sanitized_report, report_content)
        word_path = generate_word_report(report_id, report_type, sanitized_report)
        
        # Aggiorna i percorsi dei file nel database
        update_report_file_paths(report_id, pdf_path, excel_path, word_path)
        
        return report
    except json.JSONDecodeError as e:
        print(f"Errore nel parsing JSON: {str(e)}")
        print(f"Contenuto che ha causato l'errore: {content}")
        return {
            "error": f"Generazione report fallita: Errore nel parsing JSON - {str(e)}",
            "title": f"Report {report_type} - Errore Generazione",
            "executive_summary": "Impossibile generare il report automatico",
            "technical_details": f"Errore nel parsing JSON: {str(e)}\nContenuto: {content}",
            "conclusions": "Richiede generazione manuale",
            "signature": "Sistema BrokerFlow AI"
        }
    except Exception as e:
        print(f"Errore generale: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "error": f"Generazione report fallita: {str(e)}",
            "title": f"Report {report_type} - Errore Generazione",
            "executive_summary": "Impossibile generare il report automatico",
            "technical_details": str(e),
            "conclusions": "Richiede generazione manuale",
            "signature": "Sistema BrokerFlow AI"
        }

def save_compliance_report(report_type, period_start, period_end, content):
    """Salva il report di compliance nel database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO compliance_reports (report_type, period_start, period_end, content, generated_at)
        VALUES (%s, %s, %s, %s, NOW())
    """, (report_type, period_start, period_end, json.dumps(content)))
    
    report_id = cursor.lastrowid
    conn.commit()
    conn.close()
    
    return report_id

def update_report_file_paths(report_id, pdf_path, excel_path, word_path):
    """Aggiorna i percorsi dei file nel database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        UPDATE compliance_reports 
        SET file_path = %s,
            excel_path = %s,
            word_path = %s
        WHERE id = %s
    """, (pdf_path, excel_path, word_path, report_id))
    
    conn.commit()
    conn.close()

def get_compliance_reports(report_type=None):
    """Recupera i report di compliance"""
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    if report_type:
        cursor.execute("""
            SELECT * FROM compliance_reports 
            WHERE report_type = %s 
            ORDER BY generated_at DESC
        """, (report_type,))
    else:
        cursor.execute("""
            SELECT * FROM compliance_reports 
            ORDER BY generated_at DESC
        """)
    
    reports = cursor.fetchall()
    conn.close()
    
    # Converti gli oggetti datetime in stringhe
    for report in reports:
        if 'generated_at' in report and report['generated_at']:
            report['generated_at'] = report['generated_at'].isoformat()
        if 'period_start' in report and report['period_start']:
            report['period_start'] = report['period_start'].isoformat()
        if 'period_end' in report and report['period_end']:
            report['period_end'] = report['period_end'].isoformat()
    
    return reports

def generate_pdf_report(report_id, report_type, report_content):
    """Genera un report PDF e salva il file"""
    try:
        # Crea directory per i report se non esiste
        reports_dir = os.path.join(os.getcwd(), "output", "compliance_reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        # Genera nome file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type}_report_{report_id}_{timestamp}.pdf"
        file_path = os.path.join(reports_dir, filename)
        
        # Crea il documento PDF
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Titolo
        title_content = report_content.get("title", f"Report {report_type}")
        if isinstance(title_content, (dict, list)):
            title_content = str(title_content)
        title = Paragraph(title_content, styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Riepilogo esecutivo
        exec_title = Paragraph("Riepilogo Esecutivo", styles['Heading2'])
        story.append(exec_title)
        exec_summary_content = report_content.get("executive_summary", "")
        if isinstance(exec_summary_content, (dict, list)):
            exec_summary_content = str(exec_summary_content)
        exec_summary = Paragraph(exec_summary_content, styles['Normal'])
        story.append(exec_summary)
        story.append(Spacer(1, 12))
        
        # Dettagli tecnici
        tech_title = Paragraph("Dettagli Tecnici", styles['Heading2'])
        story.append(tech_title)
        tech_details_content = report_content.get("technical_details", "")
        # Converti in stringa se è un dizionario o lista
        if isinstance(tech_details_content, (dict, list)):
            tech_details_content = str(tech_details_content)
        tech_details = Paragraph(tech_details_content, styles['Normal'])
        story.append(tech_details)
        story.append(Spacer(1, 12))
        
        # Conclusioni
        concl_title = Paragraph("Conclusioni", styles['Heading2'])
        story.append(concl_title)
        conclusions_content = report_content.get("conclusions", "")
        if isinstance(conclusions_content, (dict, list)):
            conclusions_content = str(conclusions_content)
        conclusions = Paragraph(conclusions_content, styles['Normal'])
        story.append(conclusions)
        story.append(Spacer(1, 12))
        
        # Firma
        signature_content = report_content.get("signature", "Sistema BrokerFlow AI")
        if isinstance(signature_content, (dict, list)):
            signature_content = str(signature_content)
        signature = Paragraph(f"<b>Firma:</b> {signature_content}", styles['Normal'])
        story.append(signature)
        
        # Costruisci il PDF
        doc.build(story)
        
        return file_path
    except Exception as e:
        print(f"Errore nella generazione del PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

def generate_excel_report(report_id, report_type, report_content, raw_data):
    """Genera un report Excel e salva il file"""
    try:
        # Crea directory per i report se non esiste
        reports_dir = os.path.join(os.getcwd(), "output", "compliance_reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        # Genera nome file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type}_report_{report_id}_{timestamp}.xlsx"
        file_path = os.path.join(reports_dir, filename)
        
        # Crea un DataFrame pandas con i dati
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Foglio 1: Riepilogo
            summary_data = {
                'Sezione': ['Titolo', 'Tipo Report', 'Periodo Inizio', 'Periodo Fine', 'Riepilogo Esecutivo', 'Conclusioni', 'Firma'],
                'Contenuto': [
                    str(report_content.get("title", "")),
                    report_type,
                    str(report_content.get("period_start", "")),
                    str(report_content.get("period_end", "")),
                    str(report_content.get("executive_summary", "")),
                    str(report_content.get("conclusions", "")),
                    str(report_content.get("signature", ""))
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Riepilogo', index=False)
            
            # Foglio 2: Dettagli Tecnici
            tech_data = {
                'Dettagli': [str(report_content.get("technical_details", ""))]
            }
            tech_df = pd.DataFrame(tech_data)
            tech_df.to_excel(writer, sheet_name='Dettagli Tecnici', index=False)
            
            # Foglio 3: Dati Raw (se disponibili)
            print(f"Tipo di report: {report_type}")
            print(f"Dati raw disponibili: {list(raw_data.keys()) if raw_data else 'Nessun dato'}")
            
            if report_type == "GDPR" and "data_processing_activities" in raw_data:
                print(f"Dati attività di processamento: {raw_data['data_processing_activities']}")
                activities_df = pd.DataFrame(raw_data["data_processing_activities"])
                activities_df.to_excel(writer, sheet_name='Attività di Processamento', index=False)
            elif report_type == "SOX" and "financial_summary" in raw_data:
                print(f"Dati finanziari: {raw_data['financial_summary']}")
                financial_df = pd.DataFrame([raw_data["financial_summary"]])
                financial_df.to_excel(writer, sheet_name='Dati Finanziari', index=False)
            elif report_type == "IVASS" and "portfolio_analysis" in raw_data:
                print(f"Analisi portafoglio: {raw_data['portfolio_analysis']}")
                portfolio_df = pd.DataFrame(raw_data["portfolio_analysis"])
                portfolio_df.to_excel(writer, sheet_name='Analisi Portafoglio', index=False)
        
        return file_path
    except Exception as e:
        print(f"Errore nella generazione dell'Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

def generate_word_report(report_id, report_type, report_content):
    """Genera un report Word e salva il file"""
    try:
        # Crea directory per i report se non esiste
        reports_dir = os.path.join(os.getcwd(), "output", "compliance_reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        # Genera nome file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type}_report_{report_id}_{timestamp}.docx"
        file_path = os.path.join(reports_dir, filename)
        
        # Crea il documento Word
        doc = Document()
        
        # Titolo
        doc.add_heading(str(report_content.get("title", f"Report {report_type}")), 0)
        
        # Riepilogo esecutivo
        doc.add_heading('Riepilogo Esecutivo', level=1)
        doc.add_paragraph(str(report_content.get("executive_summary", "")))
        
        # Dettagli tecnici
        doc.add_heading('Dettagli Tecnici', level=1)
        doc.add_paragraph(str(report_content.get("technical_details", "")))
        
        # Conclusioni
        doc.add_heading('Conclusioni', level=1)
        doc.add_paragraph(str(report_content.get("conclusions", "")))
        
        # Firma
        doc.add_heading('Firma', level=1)
        doc.add_paragraph(str(report_content.get("signature", "Sistema BrokerFlow AI")))
        
        # Salva il documento
        doc.save(file_path)
        
        return file_path
    except Exception as e:
        print(f"Errore nella generazione del Word: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

def send_report_via_email(report_id, recipient_email, format_type="pdf"):
    """Invia un report via email"""
    try:
        # Recupera il report dal database
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT * FROM compliance_reports WHERE id = %s", (report_id,))
        report = cursor.fetchone()
        
        conn.close()
        
        if not report:
            return {"success": False, "error": "Report non trovato"}
        
        # Determina il percorso del file in base al formato
        file_path = None
        if format_type.lower() == "pdf":
            file_path = report.get('file_path')
            mime_type = 'application/pdf'
        elif format_type.lower() == "excel":
            file_path = report.get('excel_path')
            mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type.lower() == "word":
            file_path = report.get('word_path')
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return {"success": False, "error": "Formato non supportato"}
        
        if not file_path or not os.path.exists(file_path):
            return {"success": False, "error": "File del report non trovato"}
        
        # Configura l'email
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = recipient_email
        msg['Subject'] = f"Report di Compliance - {report.get('report_type', 'N/A')}"
        
        # Corpo dell'email
        body = f"""
        Gentile destinatario,
        
        In allegato trova il report di compliance per il periodo {report.get('period_start', '')} - {report.get('period_end', '')}.
        
        Tipo report: {report.get('report_type', 'N/A')}
        
        Cordiali saluti,
        Team BrokerFlow AI
        """
        msg.attach(MIMEText(body, 'plain'))
        
        # Allega il file
        with open(file_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
        
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename= {os.path.basename(file_path)}'
        )
        msg.attach(part)
        
        # Invia l'email
        if EMAIL_CONFIG['smtp_port'] == 465:
            # Usa SMTP_SSL per la porta 465
            server = smtplib.SMTP_SSL(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        else:
            # Usa SMTP con STARTTLS per altre porte
            server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
            server.starttls()
        
        server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
        text = msg.as_string()
        server.sendmail(EMAIL_CONFIG['sender_email'], recipient_email, text)
        server.quit()
        
        # Salva il log dell'invio
        save_email_log(report_id, recipient_email, format_type)
        
        return {"success": True, "message": "Email inviata con successo"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def save_email_log(report_id, recipient_email, format_type):
    """Salva il log dell'invio email"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO email_logs (report_id, recipient_email, format_type, sent_at)
        VALUES (%s, %s, %s, NOW())
    """, (report_id, recipient_email, format_type))
    
    conn.commit()
    conn.close()